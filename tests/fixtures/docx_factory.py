"""테스트용 DOCX 파일 생성 헬퍼."""

from __future__ import annotations

import io

from docx import Document


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """주어진 단락·표 내용으로 DOCX bytes를 생성한다."""
    doc = Document()

    for text in paragraphs:
        doc.add_paragraph(text)

    if table_rows:
        cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=0, cols=cols)
        for row_data in table_rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                if i < cols:
                    row.cells[i].text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_minimal_pdf_bytes() -> bytes:
    """텍스트 없는 최소 유효 PDF bytes (파이프라인 테스트용)."""
    return (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n"
        b"xref\n"
        b"0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer<</Size 4/Root 1 0 R>>\n"
        b"startxref\n"
        b"178\n"
        b"%%EOF"
    )
