"""자료 파싱 + MaterialIngestor — Step 6 스텁 구현.

사용자가 업로드한 자료 파일(PDF·DOCX·TXT)을 텍스트로 추출하고
PII 마스킹을 거쳐 MaterialDoc/MaterialBundle을 반환한다.

보안 원칙: PII masker는 반드시 LLM 호출 전에 실행한다.
이 모듈이 반환하는 MaterialDoc.masked_text만 LLM에 전달할 것.

LLM 기반 문서 요약(summary 필드)은 Step 11에서 구현한다.
현재는 masked_text 앞부분을 잘라 stub 요약으로 사용한다.
"""

from __future__ import annotations

import io
import logging

from app.models import MaterialDoc, MaterialBundle
from app.pii.regex_masker import mask_text

logger = logging.getLogger(__name__)

_SUMMARY_STUB_LENGTH = 300   # 스텁 요약: 앞 300자
_SUPPORTED_EXTS = frozenset(["pdf", "docx", "doc", "txt", "md"])


# ──────────────────────────────────────────────
# 공개 API
# ──────────────────────────────────────────────

def ingest_file(file_bytes: bytes, filename: str) -> MaterialDoc:
    """파일 bytes를 파싱해 PII 마스킹된 MaterialDoc을 반환한다.

    Args:
        file_bytes: 업로드된 파일 원본 bytes
        filename:   원본 파일명 (확장자 추출에 사용)

    Returns:
        MaterialDoc — masked_text는 PII 마스킹 완료. summary는 스텁(Step 11에서 교체).

    Raises:
        ValueError: 지원하지 않는 파일 형식
    """
    ext = _get_ext(filename)
    if ext not in _SUPPORTED_EXTS:
        raise ValueError(f"지원하지 않는 파일 형식입니다: .{ext} (지원: {sorted(_SUPPORTED_EXTS)})")

    raw_text = _extract_text(file_bytes, ext, filename)

    # ── PII 마스킹 (LLM 호출 전 필수) ──────────
    mask_result = mask_text(raw_text)
    if mask_result.has_pii:
        logger.info("[PII] '%s'에서 %d개 PII 마스킹 완료", filename, len(mask_result.mask_map))

    # ── 스텁 요약 (Step 11에서 LLM 요약으로 교체) ──
    summary = _stub_summary(mask_result.masked_text)

    return MaterialDoc(
        source_name=filename,
        masked_text=mask_result.masked_text,
        summary=summary,
        doc_type=ext,
    )


def build_material_bundle(files: list[tuple[bytes, str]]) -> MaterialBundle:
    """여러 파일을 처리해 MaterialBundle을 반환한다.

    개별 파일 파싱 실패는 소프트 실패로 처리 (다른 파일 처리 계속).

    Args:
        files: [(file_bytes, filename), ...] 리스트

    Returns:
        MaterialBundle — 성공적으로 파싱된 docs 포함
    """
    docs: list[MaterialDoc] = []
    for file_bytes, filename in files:
        try:
            doc = ingest_file(file_bytes, filename)
            docs.append(doc)
        except Exception as exc:
            logger.warning("[MaterialIngestor] '%s' 파싱 실패 (건너뜀): %s", filename, exc)
    return MaterialBundle(docs=docs)


# ──────────────────────────────────────────────
# 파일 형식별 텍스트 추출
# ──────────────────────────────────────────────

def _extract_text(file_bytes: bytes, ext: str, filename: str) -> str:
    """확장자에 따라 적절한 파서를 호출한다."""
    if ext == "pdf":
        return _extract_pdf(file_bytes, filename)
    if ext in ("docx", "doc"):
        return _extract_docx(file_bytes, filename)
    if ext in ("txt", "md"):
        return _extract_plain(file_bytes)
    return ""


def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    """PDF에서 텍스트를 추출한다.

    1차: pdfplumber (표·레이아웃 인식 우수)
    2차 fallback: pypdf (pdfplumber 실패 시)
    """
    text_parts: list[str] = []

    # ── pdfplumber ──────────────────────────────
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
                # 표도 텍스트로 변환
                for table in page.extract_tables() or []:
                    for row in table:
                        row_text = " | ".join(cell or "" for cell in row if cell)
                        if row_text.strip():
                            text_parts.append(row_text)
        if text_parts:
            return "\n".join(text_parts)
    except Exception as exc:
        logger.debug("[PDF] pdfplumber 실패, pypdf로 재시도: %s", exc)

    # ── pypdf fallback ──────────────────────────
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as exc:
        logger.warning("[PDF] '%s' 텍스트 추출 실패: %s", filename, exc)
        return ""


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    """DOCX에서 텍스트를 추출한다 (단락 + 표)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        # 단락
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # 표
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)
    except Exception as exc:
        logger.warning("[DOCX] '%s' 텍스트 추출 실패: %s", filename, exc)
        return ""


def _extract_plain(file_bytes: bytes) -> str:
    """TXT/MD 파일을 UTF-8로 디코딩한다."""
    for encoding in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


# ──────────────────────────────────────────────
# 내부 유틸
# ──────────────────────────────────────────────

def _get_ext(filename: str) -> str:
    """파일명에서 소문자 확장자를 반환한다."""
    if "." not in filename:
        return "txt"
    return filename.rsplit(".", 1)[-1].lower()


def _stub_summary(text: str) -> str:
    """Step 11 이전의 임시 요약: 앞 300자 발췌."""
    if not text.strip():
        return ""
    truncated = text.strip()[:_SUMMARY_STUB_LENGTH]
    if len(text.strip()) > _SUMMARY_STUB_LENGTH:
        truncated += "…"
    return f"[요약 미생성 — Step 11에서 LLM 요약으로 교체 예정]\n{truncated}"
